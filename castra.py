import numpy as np
from PyQt5.QtCore import QDir, QTimer  # Import QDir and QTimer
from PyQt5.QtWidgets import QDialog, QVBoxLayout
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanva
import matplotlib.pyplot as plt
from pymatgen.analysis.diffraction.xrd import XRDCalculator
from pymatgen.io.cif import CifParser
import sys
import math
from PySide2.QtGui import QPixmap  # Correct import statement
import re
import csv
from PyQt5.QtCore import Qt  # Add this imports
import os
import pandas as pd
from PySide2.QtGui import QIcon, QFont
from PySide2.QtWidgets import QSizePolicy,QSplitter,QApplication, QMainWindow, QAction, QTextEdit, QVBoxLayout,QFileSystemModel,  QTreeView, QHBoxLayout, QWidget, QFileDialog, QInputDialog, QPushButton, QComboBox, QLabel, QLineEdit, QDialog, QTableWidget, QTableWidgetItem, QMessageBox
from docx import Document  # Import the Document class for creating Word documents
import numpy as np
import matplotlib.pyplot as plt
from pymatgen.analysis.diffraction.xrd import XRDCalculator
from pymatgen.io.cif import CifParser
#from PyQt5.QtCore import QDir, QTimer  # Import QDir and QTimer
from PySide2.QtCore import  QDir

class CIFFileViewer(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Dimerc PairsViewer")
        self.setGeometry(50, 50, 1200, 600)

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        layout = QVBoxLayout(self.central_widget)

        self.model = QFileSystemModel()
        self.model.setRootPath(os.getcwd())  # Set the root path to the directory containing MLC file

        self.model.setFilter(QDir.Files | QDir.NoDotAndDotDot)
        self.model.setNameFilters(["trans*"])  # Only files with .cif extension
        self.model.setNameFilterDisables(False)

        self.tree_view = QTreeView()
       
        
        self.tree_view.setModel(self.model)
        self.tree_view.setRootIndex(self.model.index(os.getcwd()))  # Set the root index to the directory containing MLC file
        layout.addWidget(self.tree_view)
        self.text_edit = QTextEdit()
        layout.addWidget(self.text_edit)
        #self.table_widget = QTableWidget()
        #self.table_widget.setColumnCount(8)  # Adjust number of columns as needed
        #self.table_widget.setHorizontalHeaderLabels(["A1","A2","Distance","Columbic","Polarization","Dispersion","Repulsion","Total"])  # Adjust column headers as needed
        
        self.timer = QTimer()
        self.timer.setInterval(2000)  # Set interval to 2000 milliseconds (2 seconds)
        self.timer.setSingleShot(True)  # Set single shot mode

        self.timer.timeout.connect(self.open_next_cif)

        self.tree_view.clicked.connect(self.start_timer)
    
    
    def calculate_bond_distance(self,atom1, atom2):
        return math.sqrt((atom1[0] - atom2[0]) ** 2 + (atom1[1] - atom2[1]) ** 2 + (atom1[2] - atom2[2]) ** 2)

    def calculate_bond_angle(self,atom1, atom2, atom3):
        vector1 = (atom1[0] - atom2[0], atom1[1] - atom2[1], atom1[2] - atom2[2])
        vector2 = (atom3[0] - atom2[0], atom3[1] - atom2[1], atom3[2] - atom2[2])
        dot_product = sum(x * y for x, y in zip(vector1, vector2))
        magnitude1 = math.sqrt(sum(x ** 2 for x in vector1))
        magnitude2 = math.sqrt(sum(x ** 2 for x in vector2))
        angle_rad = math.acos(dot_product / (magnitude1 * magnitude2))
        return math.degrees(angle_rad)

    def check_hydrogen_bonds(self,molecule1_atoms, molecule2_atoms):
        hydrogen_bonds = []
        HYDROGEN_BOND_DISTANCE = 3.0  
        HYDROGEN_BOND_ANGLE = 120.0  
        
        for atom1 in molecule1_atoms:
            if atom1[0].startswith('H'):
                for atom2 in molecule2_atoms:
                    if atom2[0].startswith('O') or atom2[0].startswith('N'):
                        distance = self.calculate_bond_distance(atom1[1:], atom2[1:])
                        if distance <= HYDROGEN_BOND_DISTANCE:
                            angle = self.calculate_bond_angle(atom1[1:], atom2[1:], molecule2_atoms[0][1:])
                            if angle >= HYDROGEN_BOND_ANGLE:
                                hydrogen_bonds.append((atom1[0], atom2[0], distance, angle))
        return hydrogen_bonds
    
    def check_halogen_bonds(self,molecule1_atoms, molecule2_atoms):
        halogen_bonds = []
        HALOGEN_BOND_DISTANCE = 4.0  
        HALOGEN_BOND_ANGLE = 150.0  

        for atom1 in molecule1_atoms:
            if atom1[0].startswith('F') or atom1[0].startswith('Cl') or atom1[0].startswith('Br') or atom1[0].startswith('I'):
                for atom2 in molecule2_atoms:
                    if atom2[0].startswith('C') or atom2[0].startswith('N') or atom2[0].startswith('O') or atom2[0].startswith('S'):
                        distance = self.calculate_bond_distance(atom1[1:], atom2[1:])
                        if distance <= HALOGEN_BOND_DISTANCE:
                            angle = self.calculate_bond_angle(atom1[1:], atom2[1:], molecule2_atoms[0][1:])
                            if angle >= HALOGEN_BOND_ANGLE:
                                halogen_bonds.append((atom1[0], atom2[0], distance, angle))
        return halogen_bonds

    
    def calculate_vdw_interaction(atom1, atom2):

        r = calculate_bond_distance(atom1[1:], atom2[1:])
    # Lennard-Jones potential constants
        epsilon = (atom1[3] + atom2[3]) / 2  # average depth of the potential well
        sigma = (atom1[4] + atom2[4]) / 2  # effective diameter
        if r == 0:
            return 0
        return 4 * epsilon * ((sigma / r)**12 - (sigma / r)**6)
        
    def calculate_electrostatic_interaction(atom1, atom2):

        k = 8.99 * 10**9  # Coulomb's constant in N*m^2/C^2
        q1 = atom1[2]  # partial charge on atom1
        q2 = atom2[2]  # partial charge on atom2
        r= calculate_bond_distance(atom1[1:], atom2[1:])
        if r == 0:
            return 0
        return k * q1 * q2 / r**2
        
        
    def parse_molecule_data(self,file_path):
        molecule1 = []
        molecule2 = []
        
        current_molecule = molecule1
       
        with open(file_path, 'r') as file:
                for line in file:
                    if line.strip() and line[0].isalpha():
                        match = re.match(r'([A-Za-z0-9]+)\s+[A-Za-z]+\s+([\d.-]+)\s+([\d.-]+)\s+([\d.-]+)', line)
                        if match:
                            atom_label = match.group(1)
                            atom_coords = tuple(map(float, match.group(2, 3, 4)))
                            if any(atom[0] == atom_label for atom in molecule1):
                                current_molecule = molecule2
                            current_molecule.append((atom_label, *atom_coords))

        return molecule1, molecule2
    
    def start_timer(self, index):
        self.timer.start()

    def open_cif(self, index):
        #print("hai called ",index)
        file_path = self.model.filePath(index)
        fcont=open(file_path,"r")
        content=fcont.readlines()
        
        self.text_edit.setPlainText(''.join(content))
       
        molecule1, molecule2 = self.parse_molecule_data(file_path)
        
                
            
        hydrogen_bonds = self.check_hydrogen_bonds(molecule1, molecule2)
        halogen_bonds = self.check_halogen_bonds(molecule1, molecule2)
        
        hydrogen_bonds_df = pd.DataFrame(hydrogen_bonds, columns=['Atom 1', 'Atom 2', 'Distance (Å)', 'Angle (°)'])
        halogen_bonds_df = pd.DataFrame(halogen_bonds, columns=['Atom 1', 'Atom 2', 'Distance (Å)', 'Angle (°)'])

        #print("Hydrogen Bonds:")
        hbs=hydrogen_bonds_df.to_string(index=False)

        habs=halogen_bonds_df.to_string(index=False)
        
        self.text_edit.setPlainText(hbs+"\n"+habs)
        

        #print("\nHalogen Bonds:")
        #print(halogen_bonds_df)

         
    def open_next_cif(self):
        index = self.tree_view.currentIndex()

        next_index = index.sibling(index.row() , index.column())

        if next_index.isValid():
            self.open_cif(next_index)
        
        
class AddTrans(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.initUI()
    def initUI(self):
        self.setWindowTitle("Translation Operation")
        layout = QVBoxLayout()
        self.normEdit = QLineEdit()
        layout.addWidget(QLabel("Translation Vector"))
        layout.addWidget(self.normEdit)

        self.saveButton = QPushButton("Save")
        self.saveButton.clicked.connect(self.saveDataT)
        layout.addWidget(self.saveButton)

        self.setLayout(layout)

    def saveDataT(self):
        atom_number = self.normEdit.text()
        data = f"{atom_number}"
        self.parent().addTrans(data)
        self.close()


class AddSymmetryNormalVector(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.initUI()
    def initUI(self):
        self.setWindowTitle("Symmetry Operation using Normal Vector ")
        layout = QVBoxLayout()

      
        self.normEdit = QLineEdit()
        layout.addWidget(QLabel("Please input the coefficients of the normal vector of the mirror plane (e.g., for x + y = 0, enter 1 1 0):"))
        layout.addWidget(self.normEdit)

        self.saveButton = QPushButton("Save")
        self.saveButton.clicked.connect(self.saveDataS)
        layout.addWidget(self.saveButton)

        self.setLayout(layout)

    def saveDataS(self):
        atom_number = self.normEdit.text()
        data = f"{atom_number}"
        self.parent().addSymm(data)
        self.close()


class AddAtomsDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.initUI()

    def initUI(self):
        self.setWindowTitle("Atoms Module ")
        layout = QVBoxLayout()

        self.atomComboBox = QComboBox()
        self.atomComboBox.addItems(['C', 'H', 'Br', 'N', 'O', 'Cl', 'F'])
        layout.addWidget(self.atomComboBox)

        self.atomNumberEdit = QLineEdit()
        layout.addWidget(QLabel("Atom Number:"))
        layout.addWidget(self.atomNumberEdit)

        coordinatesLayout = QHBoxLayout()
        self.xCoordEdit = QLineEdit()
        self.yCoordEdit = QLineEdit()
        self.zCoordEdit = QLineEdit()
        coordinatesLayout.addWidget(QLabel("X:"))
        coordinatesLayout.addWidget(self.xCoordEdit)
        coordinatesLayout.addWidget(QLabel("Y:"))
        coordinatesLayout.addWidget(self.yCoordEdit)
        coordinatesLayout.addWidget(QLabel("Z:"))
        coordinatesLayout.addWidget(self.zCoordEdit)
        layout.addLayout(coordinatesLayout)

        self.saveButton = QPushButton("Save")
        self.saveButton.clicked.connect(self.saveDataend)
        layout.addWidget(self.saveButton)

        self.setLayout(layout)
    

    def saveData(self):
        atom = self.atomComboBox.currentText()
        atom_number = self.atomNumberEdit.text()
        x = self.xCoordEdit.text()
        y = self.yCoordEdit.text()
        z = self.zCoordEdit.text()
       
        data = f"{atom} {atom}{atom_number}  {x} {y} {z}"
        self.parent().addAtomsBeginning(data)
        self.close()
        
    def saveDataend(self):
        atom = self.atomComboBox.currentText()
        atom_number = self.atomNumberEdit.text()
        x = self.xCoordEdit.text()
        y = self.yCoordEdit.text()
        z = self.zCoordEdit.text()
       
        data = f"{atom} {atom}{atom_number}  {x} {y} {z}"
        self.parent().addAtoms(data)
        self.close()

class CIFEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('CASTRA - Crystallographic analysis, Symmetry , Transformation , Representation and Automation -Version1.0')
        self.setGeometry(40, 40, 1300, 650)
        self.setStyleSheet("""
        QMainWindow {
            background-color: rgba(240, 240, 240, 255); 
        }
        QMenuBar {
            background-color: #333;
            color: white;
        }
        QMenuBar::item {
            background: #333;
            color: white;
        }
        QMenuBar::item:selected {
            background: #555;
        }
        QToolBar {
            background-color: #ddd;
        }
        QAction {
            font-size: 12px;
        }
    """)

        self.createMenu()
        self.createMainWidget()

    def createMenu(self):
        menubar = self.menuBar()

        fileMenu = menubar.addMenu('&File')

        openAction = QAction(QIcon('open.png'), '&Open', self)
        openAction.setShortcut('Ctrl+O')
        openAction.setStatusTip('Open CIF file')
        openAction.triggered.connect(self.openCIF)
        fileMenu.addAction(openAction)

        saveAction = QAction(QIcon('save.png'), '&Save', self)
        saveAction.setShortcut('Ctrl+S')
        saveAction.setStatusTip('Save CIF file')
        saveAction.triggered.connect(self.saveCIF)
        fileMenu.addAction(saveAction)

        editMenu = menubar.addMenu('&Edit')

        addEndAction = QAction('Add Atoms at End', self)
        addEndAction.setStatusTip('Add atoms at the end of the file')
        addEndAction.triggered.connect(self.openAddAtomsDialogEnd)
        editMenu.addAction(addEndAction)
        
        
        symaddEndAction = QAction('Symmetry operations using Normal Vector of the Mirror Plane', self)
        symaddEndAction.setStatusTip('Symmetry operations using Normal Vector')
        symaddEndAction.triggered.connect(self.openAddSymDialog)
        editMenu.addAction(symaddEndAction)
        
        transaddEndAction = QAction('Translate the coordinates by a Vector ', self)
        transaddEndAction.setStatusTip('Translate the coordinates by a Vector')
        transaddEndAction.triggered.connect(self.openAddTransDialog)
        editMenu.addAction(transaddEndAction)
        
        deleteAtomAction = QAction('Delete Atom', self)
        deleteAtomAction.setStatusTip('Delete a particular atom')
        deleteAtomAction.triggered.connect(self.deleteAtom)
        editMenu.addAction(deleteAtomAction)

        editAtomAction = QAction('Edit Atom', self)
        editAtomAction.setStatusTip('Edit a particular atom')
        editAtomAction.triggered.connect(self.editAtom)
        editMenu.addAction(editAtomAction)

        cifCalculationsMenu = menubar.addMenu('&CIF Calculations')

        bondDistanceAction = QAction('Bond Distance Calculation', self)
        bondDistanceAction.setStatusTip('Calculate bond distances')
        bondDistanceAction.triggered.connect(self.calculateBondDistance)
        cifCalculationsMenu.addAction(bondDistanceAction)

        bondAngleAction = QAction('Bond Angle Calculation', self)
        bondAngleAction.setStatusTip('Calculate bond angles')
        bondAngleAction.triggered.connect(self.calculateBondAngle)
        cifCalculationsMenu.addAction(bondAngleAction)

        torsionAngleAction = QAction('Torsion Angle Calculation', self)
        torsionAngleAction.setStatusTip('Calculate torsion angles')
        torsionAngleAction.triggered.connect(self.calculateTorsionAngle)
        cifCalculationsMenu.addAction(torsionAngleAction)
        
        enrichmentRatioAction = QAction('Enrichment Ratio Calculation', self)
        enrichmentRatioAction.setStatusTip('Perform Enrichment Ratio Calculation')
        enrichmentRatioAction.triggered.connect(self.calculateEnrichmentRatio)
        cifCalculationsMenu.addAction(enrichmentRatioAction)
        torcifsdfAction = QAction('Torsion Angle - CIF and SDF ', self)
        torcifsdfAction.setStatusTip('Perform Torsion Angle calculation between SDF and CIF')
        torcifsdfAction.triggered.connect(self.calculatetorsionsdf)
        cifCalculationsMenu.addAction(torcifsdfAction)
   
        dimericPairsMenu = menubar.addMenu('Generate Dimers')

        
        singleMoleculeAction = QAction('Single Molecule-Load MLC file ', self)
        singleMoleculeAction.setStatusTip('Generate Dimeric Pair with Single Molecule')
        singleMoleculeAction.triggered.connect(self.generateDimericPairsSingleMolecule)
        dimericPairsMenu.addAction(singleMoleculeAction)

        twoMoleculesAction = QAction('Two Molecules-Load MLC file', self)
        twoMoleculesAction.setStatusTip('Generate Dimeric Pair with Two Molecules')
        twoMoleculesAction.triggered.connect(self.generateDimericPairsTwoMolecules)
        dimericPairsMenu.addAction(twoMoleculesAction)

        hbondCharacterizationAction = QAction('Hbond Characterization', self)
        hbondCharacterizationAction.setStatusTip('Characterize Hbond')
        hbondCharacterizationAction.triggered.connect(self.characterizeHbond)
        cifCalculationsMenu.addAction(hbondCharacterizationAction)

        aimTableAction = QAction('AIM Table Calculations', self)
        aimTableAction.setStatusTip('Perform AIM Table Calculations')
        aimTableAction.triggered.connect(self.performAIMTableCalculations)
        cifCalculationsMenu.addAction(aimTableAction)


        
           
    def downloadBondDistanceReport(self):
        self.calculateBondDistance()

        report = self.generateBondDistanceReport()

        self.saveReportAsWord(report)
        
    def generateBondDistanceReport(self):
        report = Document()
        report.add_heading('Bond Distance Calculation Report', level=1)
    
        report.add_heading('Bond Distances', level=2)

        table = report.add_table(rows=self.bondDistanceTable.rowCount(), cols=2)
        for i in range(len(self.bondDistanceTable)):
            bond, distance = self.bondDistanceTable.item(i, 0).text(), self.bondDistanceTable.item(i, 1).text()
            table.cell(i, 0).text = bond
            table.cell(i, 1).text = distance

        return report
        
    

    def saveReportAsWord(self, report):
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Bond Distance Report", "", "Word Files (*.docx)")
        if file_path:
            report.save(file_path)
            QMessageBox.information(self, "Success", "Report saved successfully.")

            folder_path = os.path.dirname(file_path)
            QDesktopServices.openUrl(QUrl.fromLocalFile(folder_path))
    def createMainWidget(self):
        self.mainWidget = QWidget()
        layout = QVBoxLayout()
        leftWidget = QWidget()
        leftLayout = QVBoxLayout()

        self.textEdit = QTextEdit()
        self.textEdit.setStyleSheet("background-color: #f0f0f0;")
        leftLayout.addWidget(self.textEdit)
        leftWidget.setLayout(leftLayout)
        leftWidget.setSizePolicy(QSizePolicy.Minimum, QSizePolicy.Preferred)  # Set size policy
        rightWidget = QWidget()
        rightLayout = QVBoxLayout()
        buttonsLayout = QHBoxLayout()
        '''
        addButton1 = QPushButton("Button 1")
        addButton2 = QPushButton("Button 2")
        addButton3 = QPushButton("Button 3")
        buttonsLayout.addWidget(addButton1)
        buttonsLayout.addWidget(addButton2)
        buttonsLayout.addWidget(addButton3)
        '''
        rightLayout.addLayout(buttonsLayout)
        self.essentialcif1 = QTableWidget()
        self.essentialcif1.setColumnCount(2)
        self.essentialcif1.setHorizontalHeaderLabels(["Data", "Data2"])
        self.essentialcif1.horizontalHeader().setVisible(False)
        
        rightLayout.addWidget(self.essentialcif1)
        self.essentialcif2 = QTextEdit()
        
        rightLayout.addWidget(self.essentialcif2)
        rightWidget.setLayout(rightLayout)

        self.bondDistanceTable = QTableWidget()
        self.bondDistanceTable.setColumnCount(2)
        self.bondDistanceTable.setHorizontalHeaderLabels(["Atoms", "Bond Distance"])
        self.bondDistanceTable.horizontalHeader().setVisible(False)
        self.bondDistanceTable.hide()
        layout.addWidget(self.bondDistanceTable)
        
        self.bondAngleTable = QTableWidget()
        self.bondAngleTable.setColumnCount(3)
        self.bondAngleTable.setHorizontalHeaderLabels(["Atom 1", "Atom 2", "Bond Angle"])
        self.bondAngleTable.hide()  # Hide the table initially
        layout.addWidget(self.bondAngleTable)
        
        self.torsionAngleTable = QTableWidget()
        self.torsionAngleTable.setColumnCount(4)
        self.torsionAngleTable.setHorizontalHeaderLabels(["Atom 1", "Atom 2", "Atom 3", "Torsion Angle"])
        self.torsionAngleTable.hide()  # Hide the table initially
        layout.addWidget(self.torsionAngleTable)
        splitter = QSplitter()
        splitter.addWidget(leftWidget)
        splitter.addWidget(rightWidget)

        layout.addWidget(splitter)
        
        self.mainWidget.setLayout(layout)
        self.setCentralWidget(self.mainWidget)
    
    def round_num(self,num, dec=0):
        return round(num, dec)
    
    def visualize_crystal_structure_tr(self,coordinates, title="Translated Structure"):
        fig = plt.figure()
        ax = fig.add_subplot(111, projection='3d')
        ax.scatter(coordinates[:, 0], coordinates[:, 1], coordinates[:, 2], c='b', marker='o')
        ax.set_xlabel('X')
        ax.set_ylabel('Y')
        ax.set_zlabel('Z')
        ax.set_title(title)
        plt.show()
    
    def visualize_crystal_structure_cr(self,coordinates, title="Crystal Structure"):
        fig = plt.figure()
        ax = fig.add_subplot(111, projection='3d')
        ax.scatter(coordinates[:, 0], coordinates[:, 1], coordinates[:, 2], c='b', marker='o')
        ax.set_xlabel('X')
        ax.set_ylabel('Y')
        ax.set_zlabel('Z')
        ax.set_title(title)
        plt.show()    
    def visualize_symmetry_3d(self,operation):
        fig = plt.figure()
        ax = fig.add_subplot(111, projection='3d')
        ax.quiver(0, 0, 0, operation[0], operation[1], operation[2])
        ax.scatter(operation[0], operation[1], operation[2], color='r')
        ax.text(operation[0], operation[1], operation[2], '   (' + str(operation[0]) + ', ' + str(operation[1]) + ', ' + str(operation[2]) + ')')
        ax.set_xlabel('X')
        ax.set_ylabel('Y')
        ax.set_zlabel('Z')
        plt.title('3D Representation of Symmetry Operation')
        plt.grid(True)
        plt.show()

    
    def addTrans(self, data):
        vector = np.array(list(map(float, data.split())))
        current_text = self.textEdit.toPlainText()
        lines = current_text.split('\n')
        ac=[]
        
        for i, line in enumerate(lines):
            
            xyz=re.findall("-?\d+\.\d+",line)
            if(len(xyz)==3):
                float_list = [float(x) for x in xyz]
                ac.append(float_list)
       
        transformed_coordinates = ac+ vector
        array_2d = np.array(ac)
        self.visualize_crystal_structure_tr(transformed_coordinates)
        
        self.visualize_crystal_structure_cr(array_2d)
        
      
        arr_string = np.array2string(transformed_coordinates)
        font = QFont()
        font.setPointSize(11)  # Set the font size to 12
        self.essentialcif2.setFont(font)
                
        self.essentialcif2.setPlainText("Transformed Co-ordinates"+"\n"+arr_string)
        
    def addSymm(self, data):
        mirror_plane = np.array(list(map(float, data.split())))
        current_text = self.textEdit.toPlainText()
        lines = current_text.split('\n')
        ac=[]
        
        for i, line in enumerate(lines):
            
            xyz=re.findall("-?\d+\.\d+",line)
            if(len(xyz)==3):
                float_list = [float(x) for x in xyz]
                ac.append(float_list)
       
        transformed_coordinates = np.dot(ac, mirror_plane)
        self.visualize_symmetry_3d(mirror_plane)
        #print(transformed_coordinates)
        arr_string = np.array2string(transformed_coordinates)
        font = QFont()
        font.setPointSize(11)  # Set the font size to 12
        self.essentialcif2.setFont(font)
                
        self.essentialcif2.setPlainText(arr_string)
    
    def addAtoms(self, data):
        current_text = self.textEdit.toPlainText()
        lines = current_text.split('\n')
        for i, line in enumerate(lines):
            if line.strip() == '#END':
                lines.insert(i, data)
                break
        new_text = '\n'.join(lines)
        self.textEdit.setPlainText(new_text)
    
    def openAddAtomsDialog(self):
        dialog = AddAtomsDialog(self)
        dialog.exec_()
        
    def openAddAtomsDialogEnd(self):
        dialog = AddAtomsDialog(self)
        dialog.exec_()
    
    def openAddSymDialog(self):
        dialog = AddSymmetryNormalVector(self)
        dialog.exec_()
    
    def treeviwofdimerfiles(self):
        dialog=TreeView()
        dialog.exec_()
    
    def openAddTransDialog(self):
        dialog = AddTrans(self)
        dialog.exec_()
    
    def addAtomsBeginning(self, data):
        current_text = self.textEdit.toPlainText()
        lines = current_text.split('\n')
        for i, line in enumerate(lines):
            if line.strip() == '#END':
                lines.insert(i, data)
                break
        new_text = '\n'.join(lines)
        self.textEdit.setPlainText(new_text)

   
    def deleteAtom(self):
        pass

    def editAtom(self):
        pass

    def parseUnitCellParameters(self,cif_data):
        cell_parameters = {}
        for line in cif_data:
            if line.startswith('_cell_length_a'):
                cell_parameters['a'] = line.split()[1]
            elif line.startswith('_cell_length_b'):
                cell_parameters['b'] = line.split()[1]
            elif line.startswith('_cell_length_c'):
                cell_parameters['c'] = line.split()[1]
            elif line.startswith('_cell_angle_alpha'):
                cell_parameters['alpha'] = line.split()[1]
            elif line.startswith('_cell_angle_beta'):
                cell_parameters['beta'] = line.split()[1]
            elif line.startswith('_cell_angle_gamma'):
                cell_parameters['gamma'] = line.split()[1]
            elif line.startswith('_symmetry_cell_setting'):
                cell_parameters['Crystal System'] = line.split()[1]
            elif line.startswith('_symmetry_space_group'):
                cell_parameters['Space Group'] = line.split()[1]+line.split()[-1]
            
         
                
        return cell_parameters
    
    def showcellparamters(self,unit_cell_parameters):   #uday
        
        self.essentialcif1.setHorizontalHeaderLabels(["Cell Parameter", "Value"])    
        row = 0
        for parameter, value in unit_cell_parameters.items():
            self.essentialcif1.insertRow(row)
            self.essentialcif1.setItem(row, 0, QTableWidgetItem(parameter))
            self.essentialcif1.setItem(row, 1, QTableWidgetItem(str(value)))
            row += 1
        
        self.essentialcif1.resizeColumnsToContents()
        self.essentialcif1.setColumnWidth(0, 250)
        self.essentialcif1.setColumnWidth(1, 250)
        self.essentialcif1.setMinimumWidth(200)  # Adjust width as needed
        self.essentialcif1.setMinimumHeight(10)  # Adjust height as neede 
        
        self.essentialcif1.show()
    
    
    def openCIF(self):
        options = QFileDialog.Options()
        fileName, _ = QFileDialog.getOpenFileName(self, "Open CIF File", "", "CIF Files (*.cif)", options=options)
        if fileName:
            with open(fileName, 'r') as file:
                cif_data = file.readlines()
                font = QFont()
                font.setPointSize(9)  # Set the font size to 12
                self.textEdit.setFont(font)
                FWE=open("CASTRA.cif","w")
                    
                for i  in range(len(cif_data)):
                    FWE.write(cif_data[i])
                FWE.close()
                self.textEdit.setPlainText(''.join(cif_data))
                uc=self.parseUnitCellParameters(cif_data)
                self.showcellparamters(uc)
                self.essentialcif2.setPlainText(None)
                

    def saveCIF(self):
        options = QFileDialog.Options()
        fileName, _ = QFileDialog.getSaveFileName(self,"Save CIF File", "", "CIF Files (*.cif)", options=options)
        if fileName:
            with open(fileName, 'w') as file:
                file.write(self.textEdit.toPlainText())
                
    
    def uniq(self,lst):
        seen = set()
        return [x for x in lst if not (x in seen or seen.add(x))] 
    
    def findint(self,first, second, row, col,ss,twod):
        for m in ss:
            ccp = re.findall(r'[A-Z][a-z]?', m[0])
            inum = re.findall(r'\d{1,3}\.?\d?', m[1])
            if (first == ccp[0] and second == ccp[1]) or (first == ccp[1] and second == ccp[0]):
                twod[row + 1][col + 1] = float(inum[0])
    
    # Dimer generation Program    
    def generateDimericPairsTwoMolecules(self):
        pass
    def showDataFrameLattice(self, df):
        dialog = QDialog(self)
        dialog.setWindowTitle("Lattice Table")
        layout = QVBoxLayout()

        table = QTableWidget()
        table.setColumnCount(len(df.columns))
        table.setRowCount(len(df))

        table.setHorizontalHeaderLabels(df.columns)

        for i, row in enumerate(df.values):
            for j, value in enumerate(row):
                item = QTableWidgetItem(str(value))
                table.setItem(i, j, item)
        
        table.setMinimumWidth(950)  # Adjust width as needed
        table.setMinimumHeight(600)  # Adjust height as neede   
        layout.addWidget(table)
        
        dialog.setLayout(layout)
        dialog.exec_()


    def generateDimericPairsSingleMolecule(self):
        input_file, _ = QFileDialog.getOpenFileName(self, "Upload MLC File", "", "MLC Files (*.mlc)")
        if input_file:
            with open(input_file, "r") as f:
                with open("CASTRAoutput.cif", "w") as o:
                    for line in f:
                        match = re.match(r'(\s+)(1)(\s+)(\d+)(\s+)(\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)', line)
                        if match:
               
                                    
                            x, y, z, m, n = map(float, match.group(8, 10, 12, 14, 16))
                            x, y, z, m, n = map(lambda val: round(val / 4.184, 1), (x, y, z, m, n))
                            
                            x_formatted = f"{x:8.1f}"
                            y_formatted = f"{y:8.1f}"
                            z_formatted = f"{z:8.1f}"
                            m_formatted = f"{m:8.1f}"
                            n_formatted = f"{n:8.1f}"
                            

                            o.write(f"{match.group(1)}{match.group(2)}{match.group(3)}{match.group(4)}{match.group(5)}{match.group(6)}")
                            #o.write("\n")
                            o.write(f"{x_formatted}{y_formatted}{z_formatted}{m_formatted}{n_formatted}{match.group(17)}\n")
                            #o.write("\n")
                        else:
                            
                            m=re.match(r'(\s+)(\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)(\s+)(-?\d+\.\d+)',line)
                                
                            if m is not None:
                                a,b,c,d,e,f = map(float, m.group(4,12, 20, 22, 24, 26))
                                l=m.group(1,2)
                                o.write(l[0]+""+l[1])
                                    
                                a=int(a)
                                b=int(b)
                                c=int(c)
                                d=round(d,1)
                                e=round(e,1)
                                f=round(f,1)
                                #print(a,b,c,d,e,f)
                                if a == 1:
                                    if d > 0:
                                        if len(str(d)) == 3:
                                            o.write("   x+" + str(d))
                                        else:
                                            o.write("     x+" + str(d))
                                    elif d < 0:
                                        if len(str(d)) == 4:
                                            o.write("   x" + str(d))
                                        else:
                                            o.write("     x" + str(d))
                                    elif d == 0:
                                        o.write("       x")
                                    
                                elif a == -1:
                                    if d > 0:
                                        if len(str(d)) == 3:
                                            o.write("  -x+" + str(d))
                                        else:
                                            o.write("    -x+" + str(d))
                                    elif d < 0:
                                        if len(str(d)) == 4:
                                            o.write("  -x" + str(d))
                                        else:
                                            o.write("    -x" + str(d))
                                    elif d == 0:
                                        o.write("      -x")

                                if b == 1:
                                    if e > 0:
                                        if len(str(e)) == 3:
                                            o.write("   y+" + str(e))
                                        else:
                                            o.write("     y+" + str(e))
                                    elif e < 0:
                                        if len(str(e)) == 4:
                                            o.write("   y" + str(e))
                                        else:
                                            o.write("     y" + str(e))
                                    elif e == 0:
                                        o.write("       y")
                                    
                                elif b == -1:
                                    if e > 0:
                                        if len(str(e)) == 3:
                                            o.write("  -y+" + str(e))
                                        else:
                                            o.write("    -y+" + str(e))
                                    elif e < 0:
                                        if len(str(e)) == 4:
                                            o.write("  -y" + str(e))
                                        else:
                                            o.write("    -y" + str(e))
                                    elif e == 0:
                                        o.write("      -y")
                                                            
                                if c == 1:
                                    if f > 0:
                                        if len(str(f)) == 3:
                                            o.write("   z+" + str(f) + "\n")
                                        else:
                                            o.write("     z+" + str(f) + "\n")
                                    elif f < 0:
                                        if len(str(f)) == 4:
                                            o.write("   z" + str(f) + "\n")
                                        else:
                                            o.write("     z" + str(f) + "\n")
                                    elif f == 0:
                                        o.write("       z" + "\n")
                                elif c == -1:
                                    if f > 0:
                                        if len(str(f)) == 3:
                                            o.write("  -z+" + str(f) + "\n")
                                        else:
                                            o.write("    -z+" + str(f) + "\n")
                                    elif f < 0:
                                        if len(str(f)) == 4:
                                            o.write("  -z" + str(f) + "\n")
                                        else:
                                            o.write("    -z" + str(f) + "\n")
                                    elif f == 0:
                                        o.write("      -z" + "\n")        

                o.close()       

            aout=open("CASTRAoutput.cif","r")
            asig=open("CASTRAsignicant.cif","w")
            lit=aout.readlines()
            for i in lit:    
                if re.search("^\s+1\s+\d+",i):
                
                    columns = i.split()
                    #print(columns)
                        
                    last_column = float(columns[-1])
                    if last_column<=-1:
                            asig.write(i)
                else:
                    asig.write(i)
            asig.close()

            sign=open("CASTRAsignicant.cif","r")
            pos=[];
            xe=[];
            ye=[];
            ze=[];
            for i in sign:
                if re.search("x",i):
                    xyz=i.split()
                    
                    xe.append(xyz[1])
                    ye.append(xyz[2])
                    ze.append(xyz[3])
                    
                if re.search("^\s+1\s+\d+",i):
                    
                    columns = i.split()
                    
                    pos.append(int(columns[1]))


            CIF=open("CASTRA.cif","r")
            crf=CIF.readlines()
            
            filec=""
            
            for i in crf:
                
                sr=i[0:4]
                if sr!="#END":
                    filec=filec+i
                  
            for i in range(len(pos)):
                variable=pos[i]
                
                variable-=1
                
                filen="trans1-"+str(pos[i])+".cif";
                
                trp=open(filen,"w")
                trp.write(filec)
                am= self.textEdit.toPlainText()
                
                for i in crf:
                
                    if re.search("^([A-Z][a-z]?\d+\s+[A-Z]\s+-?\d+)",i):
                        #print(i)
                        c=i.split();
                        
                        xp=re.findall(r"\((.*?)\)", c[2])
                        yp=re.findall(r"\((.*?)\)", c[3])
                        zp=re.findall(r"\((.*?)\)", c[4])
                        #print(c)
                        if '(' in c[2]:
                            value1 = c[2].split('(')[0]
                            x=float(value1)
                            #print("udayx",x)
                        else:
                            x=float(c[2])
                         
                        if '(' in c[3]:
                            value2 = c[3].split('(')[0]
                            y=float(value2)
                        
                        else:
                            y=float(c[3])
                            #print("udayy",y)
                        if '(' in c[4]:
                            value3 = c[4].split('(')[0]
                            z=float(value3)
                        else:
                            z=float(c[4])        
                            
                        
                        xco=eval(xe[variable])
                        yco=eval(ye[variable])
                        zco=eval(ze[variable])
                        #print(xco,yco,zco)    
                        f=re.findall("([A-Z0-9]+)",i[0:5])
                        ss=""
                        for k in f:
                            ss+=k
                        if(len(ss)>=4):
                            tcord=i[0:5]+"  "+str(xco)+" "+str(yco)+" "+str(zco)+"\n";
                            
                        elif(len(ss)>=3):
                            tcord=i[0:4]+" "+str(xco)+" "+str(yco)+" "+str(zco)+"\n";
                        
                        
                        #print(tcord)        
                        trp.write(tcord)

            lat=open("CASTRAsignicant.cif","r")
            import pandas as pd
            lt=[]
            for i in lat:
                
                if re.search("^\s+1\s+\d+\s+",i):
                    cc=i.split()
                    lt.append(cc)

            df = pd.DataFrame(lt, columns=['A1','A2','Distance','Columbic','Polarization','Dispersion','Repulsion','Total'])      
            df= df.astype({'Total': float}).sort_values(by='Total')
            #self.treeviwofdimerfiles()
            #self.createTreeView()
                    
        
            
            self.cif_viewer_window = CIFFileViewer()
    #        self.cif_viewer_window.showdff(df)
            self.cif_viewer_window.show()
            self.showDataFrameLattice(df)
            
    
    def showAlert(self):
        alert = QMessageBox()
        alert.setIcon(QMessageBox.Warning)
        alert.setText("Hello User.. Please upload CIF file first... Got to the File Menu and load the CIF File ")
        alert.setWindowTitle("Alert for uploading the CIF file")
        alert.setStandardButtons(QMessageBox.Ok)
        
        alert.exec_()
    


    def calculatetorsionsdf(self):
        current_text = self.textEdit.toPlainText()
        cifcontents = current_text.split("\n")
        
        alabels = []
        header = []
        
        for line in cifcontents:
            if re.search(r"[A-Z][a-z]?\d+[A-Z]?\s+[A-Z]\s+", line):
                atomlabels = re.findall(r"[A-Z][a-z]?\d+[A-Za-z]?", line)
                if atomlabels:
                    alabels.append(atomlabels[0])
            else:
                header.append(line)
        
        if current_text.strip() != "":
            input_file, _ = QFileDialog.getOpenFileName(self, "Upload SDF", "", "SDF Files (*.sdf)")
            
            if input_file:
                with open(input_file, "r") as freadtorsion:
                    onlysdfcor = []
                    for i in freadtorsion:
                        if re.search(r"-?\d+\.\d+\s+[A-Z]", i):
                            onlysdfcor.append(i.strip())

                cif_sdf = []
                for i in range(len(onlysdfcor)):
                    label = alabels[i] if i < len(alabels) else "X"  # Safe access
                    a = re.sub(r"[A-Z]", label, onlysdfcor[i], count=1)  # Replace first occurrence only
                    cif_sdf.append(a)
                
                header = [h for h in header if not re.search(r"END", h, re.IGNORECASE)]

                full_content = "\n".join(header + cif_sdf)
                self.calculateTorsionAnglesdfcif(full_content)
        
        else:
            self.showAlert()

           
    def calculateEnrichmentRatio(self):
        input_file, _ = QFileDialog.getOpenFileName(self, "Upload CSV File", "", "CSV Files (*.csv)")
        if input_file:

            with open(input_file, newline='') as csvfile:
                reader = csv.reader(csvfile)
                ss = [row for row in reader]
                
            list_ = []
            data = []
            for s in ss:
                cc = re.findall(r'[A-Z][a-z]?', s[0])
                inum = re.findall(r'\d{1,3}\.?\d?', s[1])
                list_.extend([cc[0], cc[1]])
                data.extend(s[0])
            cc = self.uniq(list_)
            dummy=cc

            twod = [[None for _ in range(len(cc)+1)] for _ in range(len(cc)+1)]
            for i in range(len(cc)+1):
                twod[i][0]=cc[i-1];
                twod[0][i]=cc[i-1];
            for u in range(len(cc)):
                for k in range(len(cc)):
                    self.findint(cc[k], dummy[u], k, u,ss,twod)

            org= [[None for _ in range(len(cc)+1)] for _ in range(len(cc)+1)]
            for i in range(1, len(twod)):
                org[0][i] = twod[0][i]
                org[i][0] = twod[i][0]
                for j in range(1, i + 1):
                    org[i][j] = twod[i][j]
                    
            for row_index in range(len(org)):
                for col_index in range(len(org[row_index])):
                    if org[row_index][col_index] is None:
                        org[row_index][col_index] = 0  
            

            ant = []
            ant1 = []
            surface = []
            org[0][0]="Atom";

            for j in range(1, len(twod)):
                sum1 = 0
                for k in range(1, len(twod)):
                    if twod[0][j] != twod[k][0]:
                        if twod[k][j] is None:
                            sum1 += 0
                        else:
                            sum1 += twod[k][j]
                ant.append(round(sum1,2))
            
            for j in range(1, len(twod)):
                sum2 = 0
                for k in range(1, len(twod)):
                    if twod[0][j] == twod[k][0]:
                        sum2 = twod[k][j]
                        if sum2 is not None:
                            ant1.append(round(sum2, 2))
                        else:
                            ant1.append(0)


            for i in range(1, len(twod)):
                sum1 = 0.5 * ant[i - 1] + ant1[i - 1]
                
                surface.append(round(sum1, 2))  

            squa = [(surface[i] ** 2) for i in range(len(surface))]
            good = [[None] * (len(cc) + 1) for _ in range(len(cc) + 1)]
            rc = [[None] * (len(cc) + 1) for _ in range(len(cc) + 1)]
            for i in range(1, len(twod)):
                for j in range(1, i + 1):
                    good[i][j] = twod[i][j]
                    if twod[0][j - 1] == twod[i - 1][0]:
                        rc[i][j] = round(squa[i - 1] / 100,1)
                    else:
                        rc[i][j] = round((surface[i - 1] * surface[j - 1] * 2) / 100,1)

            rock = [[None] * (len(cc) + 1) for _ in range(len(cc) + 1)]
            rc[0][0] = "RC"
            rock[0][0] = "RC"
            for i in range(1, len(twod)):
                rc[i][0] = twod[i][0]
                rc[0][i] = twod[0][i]
            
            
            for i in range(len(twod)):
                for j in range(len(twod)):
                    rcon = rc[i][j]
                    rock[i][j] = rcon
            for row_index in range(len(rc)):
                for col_index in range(len(rc[row_index])):
                    if rc[row_index][col_index] is None:
                        rc[row_index][col_index] = ''           
            enratio = [[None] * (len(cc) + 1) for _ in range(len(cc) + 1)]
            for i in range(1, len(twod)):
                for j in range(1, i + 1):
                    if twod[i][j] != 0.0:
                        if rock[i][j] is None or twod[i][j] is None :
                            rock[i][j]=''
                            twod[i][j]=''
                        elif  rock[i][j] > 0.9:
                            en = twod[i][j] / rock[i][j]
                            enrich = en
                            enratio[i][j] = round(enrich,2)
            enratio[0][0] = "ER"
            for i in range(1, len(twod)):
                enratio[i][0] = twod[i][0]
                enratio[0][i] = twod[0][i]
            
            for row_index in range(len(enratio)):
                for col_index in range(len(enratio[row_index])):
                    if enratio[row_index][col_index] is None:
                        enratio[row_index][col_index] = ''           
                        
            surface.insert(0,"Surface%")            

            labels = [row[0] for row in ss]
            sizes=[]
            for row in ss:
                
                
                if re.search("\.",row[1]):
                    d=re.findall('\d+\.\d+', row[1])
                else:
                    d=re.findall('\d+', row[1])
                    
                sz=float(d[0])
                
                sizes.append(sz)

            fig, ax = plt.subplots(figsize=(6, 6))  # Set the size of the figure
            ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
            ax.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
            plt.title('CRYSTAL Contact Distribution')
            plt.savefig('pie_chart.png')  # Save pie chart as image

            combined_data = org + [surface] + rc + enratio
            dialog = QDialog(self)
            dialog.setWindowTitle("Enrichment Ratio Calculation Result")
            
            layout = QVBoxLayout()    
            hbox = QHBoxLayout()
            font = QFont()
            font.setPointSize(11)  # Set the font size to 12
            
            table = QTableWidget()
            table.setRowCount(len(combined_data))
            table.setColumnCount(max(len(row) for row in combined_data))
            
            for i, row in enumerate(combined_data):
                for j, value in enumerate(row):
                    item = QTableWidgetItem(str(value))
                    table.setItem(i, j, item)
            table.setMinimumWidth(600)  # Adjust width as needed
            table.setMinimumHeight(600)  # Adjust height as neede    
            table.setFont(font)

            hbox.addWidget(table)
            pixmap = QPixmap('pie_chart.png')
            label = QLabel()
            label.setPixmap(pixmap)
            hbox.addWidget(label)
            
            layout.addLayout(hbox)
            
            dialog.setLayout(layout)
            dialog.exec_()
            
            
            
    def read_excel_file(self,file_path):
        lir=[]
        lit=[]

        try:
           
            df = pd.read_excel(file_path)
            df["BPL"]=df["BPL"]*0.5292 #rij
            df["BPL"]=round(df["BPL"],3)
            df["Rho"]=df["Rho"]*6.748
            df["Rho"]=round(df["Rho"],3)
            df["DelSqRho"]=df["DelSqRho"]*24.098
            df["DelSqRho"]=round(df["DelSqRho"],3)
            
            df["K"]=-(df["V"]/df["G"]) #v/g
            df["K"]=round(df["K"],5)
            df["V"]=round(df["V"],5)
            df["G"]=round(df["G"],5)
                
            df["Ellipticity"] =df["V"]+df["G"] #V+G      
            df["Ellipticity"]=round(df["Ellipticity"],5)
            
            df["Name"]=df["V"]*(-0.5)*627.529#DE
            df["Name"]=round(df["Name"],2)

            df = df[["Dimers",'BPL', 'Rho', 'DelSqRho',  'V', 'G','Ellipticity','K' ,'Name']]
            df = df.rename(columns={'Dimers': 'Dimers',
                                    'BPL': 'Rij', 
                                    'Rho': 'Rho', 
                                    'DelSqRho': 'Lap', 
                                    
                                    'V': 'V',
                                    'G': 'G',
                                    'Ellipticity': 'V+G',
                                    'K': 'V/G',
                                    'Name': 'D.E'
                                    })
            df.fillna('', inplace=True)
            return df
            
        except FileNotFoundError:
            print("File not found. Please provide a valid file path.")
            return None
        except Exception as e:
            print(f"An error occurred: {e}")
            return None 
    
    
    def characterizeHbond(self):
        import math
        from PyQt5.QtWidgets import QFileDialog, QTableWidget, QTableWidgetItem, QMessageBox, QWidget, QVBoxLayout

        # Van der Waals constants
        VANDERWAAL_CONSTANTS = {
            "H": 1.2, "C": 1.70, "O": 1.52, "N": 1.55, "He": 1.4, "F": 1.47,
            "Si": 2.10, "P": 1.80, "S": 1.80, "Cl": 1.75, "As": 1.85, "Se": 1.90,
            "Br": 1.85, "Te": 2.06, "I": 1.98, "B": 1.92
        }

        file_path, _ = QFileDialog.getOpenFileName(None, "Open Coordinate File", "", "Text Files (*.txt *.dat)")
        if not file_path:
            return

        try:
            atoms = []
            coordinates = []

            with open(file_path, 'r') as f:
                for line in f:
                    parts = line.strip().split()
                    if len(parts) == 4:
                        atom = parts[0]
                        coords = [float(x.replace("E", "e")) for x in parts[1:]]
                        atoms.append(atom)
                        coordinates.append(coords)

            results = []
            for i in range(1, len(coordinates) - 1, 3):
                atom1 = atoms[i - 1]
                atom2 = atoms[i]
                atom3 = atoms[i + 1]

                dist1 = sum((coordinates[i][j] - coordinates[i - 1][j]) ** 2 for j in range(3))
                dist2 = sum((coordinates[i + 1][j] - coordinates[i][j]) ** 2 for j in range(3))

                ed1 = math.sqrt(dist1) * 0.5292
                ed2 = math.sqrt(dist2) * 0.5292

                rd = VANDERWAAL_CONSTANTS.get(atom1, 0) - ed1
                ra = VANDERWAAL_CONSTANTS.get(atom3, 0) - ed2

                criteria = "Satisfied" if rd + ra > 0 and rd - ra > 0 else "Not Satisfied"
                stabilized = "Yes" if criteria == "Satisfied" else "No"

                results.append([
                    f"{atom1}{atom2}{atom3}",
                    criteria,
                    stabilized,
                    f"{rd:.3f}",
                    f"{ra:.3f}",
                    f"{rd - ra:.4f}",
                    f"{rd + ra:.4f}"
                ])

            # Create a child widget to display the table
            self.hbondTableWidget = QWidget()
            self.hbondTableWidget.setWindowTitle("Hydrogen Bond Distance Criteria")
            self.hbondTableWidget.resize(800, 400)

            layout = QVBoxLayout(self.hbondTableWidget)
            table = QTableWidget()
            layout.addWidget(table)

            headers = ["Atoms", "Criteria", "Stabilized by", "Rd", "Ra", "Rd - Ra", "Rd + Ra"]
            table.setColumnCount(len(headers))
            table.setRowCount(len(results))
            table.setHorizontalHeaderLabels(headers)

            for row_idx, row_data in enumerate(results):
                for col_idx, value in enumerate(row_data):
                    table.setItem(row_idx, col_idx, QTableWidgetItem(str(value)))

            table.resizeColumnsToContents()
            self.hbondTableWidget.show()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to process file: {str(e)}")

    def viewXRD(self):
    
        cifcon = self.textEdit.toPlainText()
        parser = CifParser(cifcon)
        
        structure = parser.get_structures()[0]
        xrd_calc = XRDCalculator()
        
        pattern = xrd_calc.get_pattern(structure)

        plt.plot(pattern.x, pattern.y, marker='o', linestyle='-')
        plt.xlabel('2\u03B8 (\u00b0)')
        plt.ylabel('Intensity (a.u.)')
        plt.title('XRD Pattern')
        plt.grid(True)
        plt.show()
        
    def performAIMTableCalculations(self):

        file_path, _ = QFileDialog.getOpenFileName(self, 'Open Excel File', '', 'Excel Files (*.xlsx)')
        if file_path:

            df = self.read_excel_file(file_path)
            if df is not None:

                self.showDataFrame(df)
    
    
    def showtorsionaglesintable(self,torsiona):   #uday
        dialog = QDialog(self)
        dialog.setWindowTitle("Torsion Angle Calculation Result-Gaussian Optimized structure")
        layout = QVBoxLayout()

        table = QTableWidget()
       
        table.setColumnCount(4)
        table.setRowCount(len(torsiona))

        
        labels=["Atoms","Torsion Angle of X-ray Structure","Torsion Angle of Optimized Structure","Difference"]
        table.setHorizontalHeaderLabels(labels)

        for i, row in enumerate(torsiona):
            for j, value in enumerate(row):
                item = QTableWidgetItem(str(value))
                table.setItem(i, j, item)
        table.setColumnWidth(0, 200)
        table.setColumnWidth(1, 300)
        table.setColumnWidth(2, 300)
        
        
        table.setMinimumWidth(1150)  # Adjust width as needed
        table.setMinimumHeight(600)  # Adjust height as neede   
        layout.addWidget(table)
        
        dialog.setLayout(layout)
        dialog.exec_()
    
    
    def performAIMTableCalculations(self):

        file_path, _ = QFileDialog.getOpenFileName(self, 'Open Excel File', '', 'Excel Files (*.xlsx)')
        if file_path:

            df = self.read_excel_file(file_path)
            if df is not None:

                self.showDataFrame(df)
    
   
    
    def showtorsionaglesintable1(self,torsiona):   #uday
        dialog = QDialog(self)
        dialog.setWindowTitle("Torsion Angle Calculation Result-CIF file")
        layout = QVBoxLayout()

        table = QTableWidget()
        table.setColumnCount(5)
        table.setRowCount(len(torsiona))

       
        labels=["Atom1","Atom2","Atom3","Atom4","Torsion Angle"]
        table.setHorizontalHeaderLabels(labels)

        for i, row in enumerate(torsiona):
            for j, value in enumerate(row):
                item = QTableWidgetItem(str(value))
                table.setItem(i, j, item)
        
        table.setMinimumWidth(950)  # Adjust width as needed
        table.setMinimumHeight(600)  # Adjust height as neede   
        layout.addWidget(table)
        
        dialog.setLayout(layout)
        dialog.exec_()
        
    def showDataFrame(self, df):
        dialog = QDialog(self)
        dialog.setWindowTitle("AIM Table Calculations Result")
        layout = QVBoxLayout()

        table = QTableWidget()
        table.setColumnCount(len(df.columns))
        table.setRowCount(len(df))

        table.setHorizontalHeaderLabels(df.columns)

        for i, row in enumerate(df.values):
            for j, value in enumerate(row):
                item = QTableWidgetItem(str(value))
                table.setItem(i, j, item)
        
        table.setMinimumWidth(950)  # Adjust width as needed
        table.setMinimumHeight(600)  # Adjust height as neede   
        layout.addWidget(table)
        
        dialog.setLayout(layout)
        dialog.exec_()
        
    def showbondlengthtable1(self, bd):
        dialog = QDialog(self)
        dialog.setWindowTitle("Bond Distance  Calculation Result")
        layout = QVBoxLayout()

        table = QTableWidget()
       
        table.setColumnCount(3)
        table.setRowCount(len(bd))

        
        labels=["Atoms1","Atoms2","Bond Length"]
        table.setHorizontalHeaderLabels(labels)

        for i, row in enumerate(bd):
            for j, value in enumerate(row):
                item = QTableWidgetItem(str(value))
                table.setItem(i, j, item)
        table.setColumnWidth(0, 100)
        table.setColumnWidth(1, 100)
        table.setColumnWidth(2, 100)
        
        
        table.setMinimumWidth(850)  # Adjust width as needed
        table.setMinimumHeight(600)  # Adjust height as neede   
        layout.addWidget(table)
        
        dialog.setLayout(layout)
        dialog.exec_()

    def showbondangletable1(self, bd):
        dialog = QDialog(self)
        dialog.setWindowTitle("Bond Angle Calculation Result")
        layout = QVBoxLayout()

        table = QTableWidget()
       
        table.setColumnCount(4)
        table.setRowCount(len(bd))

        
        labels=["Atoms1","Atoms2","Atoms3","Bond Angle"]
        table.setHorizontalHeaderLabels(labels)

        for i, row in enumerate(bd):
            for j, value in enumerate(row):
                item = QTableWidgetItem(str(value))
                table.setItem(i, j, item)
        table.setColumnWidth(0, 100)
        table.setColumnWidth(1, 100)
        table.setColumnWidth(2, 100)
        table.setColumnWidth(3, 100)
        
        
        
        table.setMinimumWidth(850)  # Adjust width as needed
        table.setMinimumHeight(600)  # Adjust height as neede   
        layout.addWidget(table)
        
        dialog.setLayout(layout)
        dialog.exec_()

    def calculateBondDistance(self):
        current_text = self.textEdit.toPlainText()
        self.bondDistanceTable.horizontalHeader().setVisible(True)
        self.torsionAngleTable.hide()
        self.bondAngleTable.hide()
        lines = current_text.split('\n')
        atom_positions = {}
        reading_atoms = False
        for line in lines:
            if line.strip().startswith("_atom_site"):
        
                reading_atoms = True
                continue
            if not reading_atoms:
                continue  
            if len(line.split()) >= 5:
                try:
                    atom, _, x, y, z = line.split()[0], line.split()[1], float(line.split()[2]), float(line.split()[3]), float(line.split()[4])
                    atom_positions[atom] = (x, y, z)
                except ValueError:

                    continue
            
        atoms = list(atom_positions.keys())
        distances = {}
        for i in range(len(atoms)):
            for j in range(i + 1, len(atoms)):
                atom1, atom2 = atoms[i], atoms[j]
                x1, y1, z1 = atom_positions[atom1]
                x2, y2, z2 = atom_positions[atom2]
                distance = math.sqrt((x2 - x1)**2 + (y2 - y1)**2 + (z2 - z1)**2)
                distance=round(distance,2)
                distances[(atom1,atom2,distance)] = distance

        
        self.showbondlengthtable1(distances)    
   
    def calculateBondAngle(self):
        current_text = self.textEdit.toPlainText()
        self.bondAngleTable.horizontalHeader().setVisible(True)
        self.bondDistanceTable.hide()
        self.torsionAngleTable.hide()
        lines = current_text.split('\n')
        atom_positions = {}
        reading_atoms = False
        for line in lines:
            if line.strip().startswith("_atom_site"):

                reading_atoms = True
                continue
            if not reading_atoms:
                continue  # Skip metadata lines
            if len(line.split()) >= 5:
                try:
                    atom, _, x, y, z = line.split()[0], line.split()[1], float(line.split()[2]), float(line.split()[3]), float(line.split()[4])
                    atom_positions[atom] = (x, y, z)
                except ValueError:

                    continue

        atoms = list(atom_positions.keys())
        bond_angles = {}
        for i in range(len(atoms)):
            for j in range(i + 1, len(atoms)):
                for k in range(j + 1, len(atoms)):
                    atom1, atom2, atom3 = atoms[i], atoms[j], atoms[k]
                    x1, y1, z1 = atom_positions[atom1]
                    x2, y2, z2 = atom_positions[atom2]
                    x3, y3, z3 = atom_positions[atom3]
                    vector1 = (x1 - x2, y1 - y2, z1 - z2)
                    vector2 = (x3 - x2, y3 - y2, z3 - z2)
                    dot_product = sum(v1 * v2 for v1, v2 in zip(vector1, vector2))
                    magnitude1 = math.sqrt(sum(v ** 2 for v in vector1))
                    magnitude2 = math.sqrt(sum(v ** 2 for v in vector2))
                    angle = math.acos(dot_product / (magnitude1 * magnitude2))
                    ang=round(math.degrees(angle),2)
                    bond_angles[(atom1,atom2,atom3,ang)] = ang

        self.showbondangletable1(bond_angles)



    import math

    import math

    def calculateTorsionAnglesdfcif(self, current_text):
        self.torsionAngleTable.horizontalHeader().setVisible(True)
        self.bondDistanceTable.hide()
        self.bondAngleTable.hide()
        
        lines = current_text.split('\n')
        atom_positions = {}

        for line in lines:
            if not line.strip():
                continue
            parts = line.split()
            if len(parts) >= 4:
                try:
                    x, y, z = map(float, parts[:3])
                    atom = parts[3]
                    atom_positions[atom] = (x, y, z)
                except ValueError:
                    continue

        atoms = list(atom_positions.keys())
        torsion_angles = {}

        for i in range(len(atoms)):
            for j in range(i + 1, len(atoms)):
                for k in range(j + 1, len(atoms)):
                    for l in range(k + 1, len(atoms)):
                        atom1, atom2, atom3, atom4 = atoms[i], atoms[j], atoms[k], atoms[l]
                        x1, y1, z1 = atom_positions[atom1]
                        x2, y2, z2 = atom_positions[atom2]
                        x3, y3, z3 = atom_positions[atom3]
                        x4, y4, z4 = atom_positions[atom4]

                        vector1 = (x2 - x1, y2 - y1, z2 - z1)
                        vector2 = (x3 - x2, y3 - y2, z3 - z2)
                        vector3 = (x4 - x3, y4 - y3, z4 - z3)

                        cross_product1 = (
                            vector1[1] * vector2[2] - vector1[2] * vector2[1],
                            vector1[2] * vector2[0] - vector1[0] * vector2[2],
                            vector1[0] * vector2[1] - vector1[1] * vector2[0]
                        )
                        cross_product2 = (
                            vector2[1] * vector3[2] - vector2[2] * vector3[1],
                            vector2[2] * vector3[0] - vector2[0] * vector3[2],
                            vector2[0] * vector3[1] - vector2[1] * vector3[0]
                        )

                        dot_product = sum(x * y for x, y in zip(cross_product1, cross_product2))
                        norm1 = math.sqrt(sum(x ** 2 for x in cross_product1))
                        norm2 = math.sqrt(sum(y ** 2 for y in cross_product2))

                        if norm1 == 0 or norm2 == 0:
                            continue  # Skip invalid cases
                        
                        value = dot_product / (norm1 * norm2)
                        value = max(min(value, 1), -1)  # Clip value to avoid math domain error

                        angle = math.degrees(math.acos(value))
                        torsion_angles[(atom1, atom2, atom3, atom4)] = round(angle, 2)

        ct1 = self.calculateTorsionAngle_compare()

        d2 = {}
        for atoms_key in torsion_angles:
            try:
                op = torsion_angles[atoms_key] - ct1.get(atoms_key, 0)
                op = round(op, 2)
                d2[(atoms_key, ct1.get(atoms_key, 'N/A'), torsion_angles[atoms_key], op)] = op
            except Exception:
                pass

        self.showtorsionaglesintable(d2)

    
    def calculateTorsionAngle_compare(self):
        current_text = self.textEdit.toPlainText()

        self.torsionAngleTable.horizontalHeader().setVisible(True)
        self.bondDistanceTable.hide()
        self.bondAngleTable.hide()

        lines = current_text.split('\n')
        atom_positions = {}
        reading_atoms = False

        for line in lines:
            if line.strip().startswith("_atom_site"):
                reading_atoms = True
                continue
            if not reading_atoms:
                continue
            if len(line.split()) >= 5:
                try:
                    parts = line.split()
                    atom = parts[0]
                    x, y, z = map(float, parts[2:5])
                    atom_positions[atom] = (x, y, z)
                except ValueError:
                    continue

        atoms = list(atom_positions.keys())
        torsion_angles = {}

        for i in range(len(atoms)):
            for j in range(i + 1, len(atoms)):
                for k in range(j + 1, len(atoms)):
                    for l in range(k + 1, len(atoms)):
                        atom1, atom2, atom3, atom4 = atoms[i], atoms[j], atoms[k], atoms[l]
                        x1, y1, z1 = atom_positions[atom1]
                        x2, y2, z2 = atom_positions[atom2]
                        x3, y3, z3 = atom_positions[atom3]
                        x4, y4, z4 = atom_positions[atom4]

                        vector1 = (x2 - x1, y2 - y1, z2 - z1)
                        vector2 = (x3 - x2, y3 - y2, z3 - z2)
                        vector3 = (x4 - x3, y4 - y3, z4 - z3)
    
                        cross_product1 = (
                            vector1[1] * vector2[2] - vector1[2] * vector2[1],
                            vector1[2] * vector2[0] - vector1[0] * vector2[2],
                            vector1[0] * vector2[1] - vector1[1] * vector2[0]
                        )
                        cross_product2 = (
                            vector2[1] * vector3[2] - vector2[2] * vector3[1],
                            vector2[2] * vector3[0] - vector2[0] * vector3[2],
                            vector2[0] * vector3[1] - vector2[1] * vector3[0]
                        )

                        dot_product = sum(x * y for x, y in zip(cross_product1, cross_product2))
                        norm1 = math.sqrt(sum(x ** 2 for x in cross_product1))
                        norm2 = math.sqrt(sum(y ** 2 for y in cross_product2))

                        if norm1 == 0 or norm2 == 0:
                            continue

                        value = dot_product / (norm1 * norm2)
                        value = max(min(value, 1), -1)

                        torsion_angle = math.degrees(math.acos(value))
                        torsion_angles[(atom1, atom2, atom3, atom4)] = round(torsion_angle, 2)

        return torsion_angles


        
    def calculateTorsionAngle(self):
        current_text = self.textEdit.toPlainText()
        
        self.torsionAngleTable.horizontalHeader().setVisible(True)
        self.bondDistanceTable.hide()
        self.bondAngleTable.hide()
        lines = current_text.split('\n')
        
        atom_positions = {}
        reading_atoms = False
        for line in lines:
            if line.strip().startswith("_atom_site"):
                reading_atoms = True
                continue
            if not reading_atoms:
                continue  
            if len(line.split()) >= 5:
                try:
                    atom, _, x, y, z = line.split()[0], line.split()[1], float(line.split()[2]), float(line.split()[3]), float(line.split()[4])
                    atom_positions[atom] = (x, y, z)
                except ValueError:
                    continue

        atoms = list(atom_positions.keys())
        torsion_angles = {}
        for i in range(len(atoms)):
            for j in range(i + 1, len(atoms)):
                for k in range(j + 1, len(atoms)):
                    for l in range(k + 1, len(atoms)):
                        atom1, atom2, atom3, atom4 = atoms[i], atoms[j], atoms[k], atoms[l]
                        x1, y1, z1 = atom_positions[atom1]
                        x2, y2, z2 = atom_positions[atom2]
                        x3, y3, z3 = atom_positions[atom3]
                        x4, y4, z4 = atom_positions[atom4]

                        vector1 = (x2 - x1, y2 - y1, z2 - z1)
                        vector2 = (x3 - x2, y3 - y2, z3 - z2)
                        vector3 = (x4 - x3, y4 - y3, z4 - z3)

                        cross_product1 = (
                            vector1[1] * vector2[2] - vector1[2] * vector2[1],
                            vector1[2] * vector2[0] - vector1[0] * vector2[2],
                            vector1[0] * vector2[1] - vector1[1] * vector2[0]
                        )
                        cross_product2 = (
                            vector2[1] * vector3[2] - vector2[2] * vector3[1],
                            vector2[2] * vector3[0] - vector2[0] * vector3[2],
                            vector2[0] * vector3[1] - vector2[1] * vector3[0]
                        )

                        torsion_angle = math.degrees(math.acos(sum(x * y for x, y in zip(cross_product1, cross_product2)) /
                                                               (math.sqrt(sum(x ** 2 for x in cross_product1)) *
                                                                math.sqrt(sum(y ** 2 for y in cross_product2)))))
                        tor=round(torsion_angle,2)

                        torsion_angles[(atom1, atom2, atom3, atom4,tor)] = tor
                        
        self.showtorsionaglesintable1(torsion_angles)
        
if __name__ == '__main__':
    app = QApplication(sys.argv)
    editor = CIFEditor()
    editor.show()
    sys.exit(app.exec_())

